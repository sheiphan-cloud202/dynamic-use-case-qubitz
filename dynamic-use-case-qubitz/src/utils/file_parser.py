"""
File parsing utilities for PDF and DOCX files from S3.
"""

import os
import sys
import logging
import tempfile
from typing import Optional
from src.services.aws_clients import s3_client

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class FileParser:
    """Utility class for parsing S3 files (PDF/DOCX) with optional dependencies."""

    @staticmethod
    def _add_efs_path():
        """Ensure EFS packages are available on sys.path."""
        efs_site_packages = "/mnt/efs/envs/strands_lambda/lib/python3.11/site-packages"
        if efs_site_packages not in sys.path:
            sys.path.insert(0, efs_site_packages)
            logger.warning("✅ EFS Python packages path added to sys.path")

    @staticmethod
    def download_s3_file(s3_url: str) -> Optional[str]:
        """Download file from S3 URL to temporary location."""
        logger.warning(f"✅ File: files download_s3_file : {s3_url}")
        try:
            bucket = key = None

            if s3_url.startswith('s3://'):
                parts = s3_url.replace('s3://', '').split('/', 1)
                if len(parts) != 2:
                    logger.error(f"Cannot parse S3 URL: {s3_url}")
                    return None
                bucket, key = parts

            elif s3_url.startswith('https://') and '.s3.amazonaws.com' in s3_url:
                parts = s3_url.replace('https://', '').split('.s3.amazonaws.com/')
                if len(parts) != 2:
                    logger.error(f"Cannot parse S3 URL: {s3_url}")
                    return None
                bucket, key = parts

            else:
                logger.error(f"Invalid S3 URL format: {s3_url}")
                return None

            file_extension = os.path.splitext(key)[1]
            temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=file_extension)

            s3_client.download_file(bucket, key, temp_file.name)
            logger.warning(f"Downloaded S3 file: {s3_url} to {temp_file.name}")
            return temp_file.name

        except Exception as e:
            logger.error(f"Error downloading S3 file {s3_url}: {e}")
            return None

    @staticmethod
    def parse_pdf(file_path: str) -> Optional[str]:
        """Parse PDF file content with fallback methods."""
        FileParser._add_efs_path()

        content = ""


        try:
            import PyPDF2
            logger.warning("✅ PyPDF2 available")
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    content += page.extract_text() + "\n"

            if content.strip():
                logger.warning(f"Successfully parsed PDF with PyPDF2: {len(content)} chars")
                return content.strip()

        except Exception as e:
            logger.error(f"PyPDF2 also failed: {e}")

        try:
            import pdfplumber
            logger.warning("✅ pdfplumber available")
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        content += text + "\n"

            if content.strip():
                logger.warning(f"Successfully parsed PDF with pdfplumber: {len(content)} chars")
                return content.strip()

        except Exception as e:
            logger.warning(f"pdfplumber failed: {e}, trying PyPDF2")
    
        logger.error(f"Failed to parse PDF: {file_path}")
        return None

    @staticmethod
    def parse_docx(file_path: str) -> Optional[str]:
        """Parse DOCX file content."""
        FileParser._add_efs_path()

        try:
            import docx
            logger.warning("✅ docx available")
            doc = docx.Document(file_path)
            content = []

            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content.append(paragraph.text.strip())

            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        content.append(" | ".join(row_text))

            result = "\n".join(content)
            logger.warning(f"Successfully parsed DOCX: {len(result)} chars")
            return result

        except Exception as e:
            logger.error(f"Error parsing DOCX {file_path}: {e}")
            return None

    @staticmethod
    def parse_s3_file(s3_url: str) -> Optional[str]:
        """Parse file from S3 URL (PDF or DOCX)."""
        if not s3_url:
            return None

        temp_file_path = FileParser.download_s3_file(s3_url)
        if not temp_file_path:
            return None

        try:
            file_extension = os.path.splitext(temp_file_path)[1].lower()

            if file_extension == '.pdf':
                return FileParser.parse_pdf(temp_file_path)
            elif file_extension in ['.docx', '.doc']:
                return FileParser.parse_docx(temp_file_path)
            else:
                logger.error(f"Unsupported file type: {file_extension}")
                return None

        finally:
            try:
                if os.path.exists(temp_file_path):
                    os.unlink(temp_file_path)
            except Exception as e:
                logger.warning(f"Failed to cleanup temp file: {e}")
